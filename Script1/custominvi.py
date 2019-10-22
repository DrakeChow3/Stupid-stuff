#! python3
# Custom Invitations as Word Documents

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH #add alignment

names=[]
doc=open('guests.txt','r')
for name in doc.readlines():
    names.append(name.strip())
print(str(names))
my_doc=docx.Document()
for i in range(0,len(names)*5,5):
    #add the first paragraph
    para=my_doc.add_paragraph('It would be a pleasure to have the company of')
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.style='Heading 6'
    my_doc.paragraphs[i].runs[0].italic = True
    my_doc.paragraphs[i].runs[0].bold = True
    #add the name in the list
    para=my_doc.add_paragraph(names[int(i/5)])
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.style='Quote'
    #add paragraph
    para=my_doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.style='Heading 6'
    my_doc.paragraphs[i+2].runs[0].italic = True
    my_doc.paragraphs[i+2].runs[0].bold = True
    #add paragraph
    para=my_doc.add_paragraph('April 1st')
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.style='Quote'
    #add the final paragraph in the page
    para=my_doc.add_paragraph('at 7 o\'clock')
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.style='Heading 6'
    my_doc.paragraphs[i+4].runs[0].italic = True
    my_doc.paragraphs[i+4].runs[0].bold = True
    if i!=((len(names)-1)*5):
        my_doc.paragraphs[i+4].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
my_doc.save('twoPage.docx')